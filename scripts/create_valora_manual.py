#!/usr/bin/env python3
"""Genera il manuale operativo ValOra in formato DOCX.

Il documento usa esclusivamente contenuti e schermate locali del progetto.
La conversione e la verifica PDF vengono eseguite separatamente dal workflow.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from PIL import Image, ImageDraw, ImageFilter
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "manual"
ASSETS = OUT / "assets"
DOCX_PATH = OUT / "Manuale-operativo-ValOra.docx"
FRAMED_ASSETS = ASSETS / "_framed"
SCREENSHOT_FULL_CM = 16.7
SCREENSHOT_NARROW_CM = 10.2

NAVY = "163A5F"
TEAL = "0B7F7A"
GREEN = "38A169"
MINT = "E8F6F3"
PALE_BLUE = "EDF4FA"
PALE_GOLD = "FFF5D9"
PALE_RED = "FDECEC"
LIGHT = "F5F7F9"
MID = "D8E1E8"
TEXT = "273746"
MUTED = "607485"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(Cm(width_cm).emu / 635)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_cm: list[float]) -> None:
    """Fissa larghezza totale, griglia e colonne in DXA per Word e LibreOffice."""
    widths_dxa = [int(Cm(width).emu / 635) for width in widths_cm]
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = table._tbl.tblGrid
    new_grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        new_grid.append(col)
    table._tbl.replace(old_grid, new_grid)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table) -> None:
    """Solo righe orizzontali: niente griglia verticale pesante."""
    tbl_pr = table._tbl.tblPr
    old = tbl_pr.find(qn("w:tblBorders"))
    if old is not None:
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "insideH", "start", "end", "insideV"):
        node = OxmlElement(f"w:{edge}")
        if edge in ("top", "bottom", "insideH"):
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), "5")
            node.set(qn("w:color"), MID)
            node.set(qn("w:space"), "0")
        else:
            node.set(qn("w:val"), "nil")
        borders.append(node)
    tbl_pr.append(borders)


def set_run_shading(run, fill: str) -> None:
    r_pr = run._r.get_or_add_rPr()
    shd = r_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        r_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def keep_rows_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)
    for cell in row.cells:
        for p in cell.paragraphs:
            p.paragraph_format.keep_together = True


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.65)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in (
        ("Title", 31, NAVY, 0, 10),
        ("Subtitle", 15, TEAL, 0, 12),
        ("Heading 1", 20, NAVY, 13, 7),
        ("Heading 2", 14, TEAL, 11, 5),
        ("Heading 3", 11.5, NAVY, 8, 3),
    ):
        st = styles[name]
        st.font.name = "Arial"
        st.font.size = Pt(size)
        st.font.bold = name != "Subtitle"
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        st = styles[name]
        st.font.name = "Arial"
        st.font.size = Pt(10.2)
        st.paragraph_format.left_indent = Cm(0.55)
        st.paragraph_format.first_line_indent = Cm(-0.25)
        st.paragraph_format.space_after = Pt(3)

    caption = styles["Caption"]
    caption.font.name = "Arial"
    caption.font.size = Pt(8.5)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = False

    if "Small note" not in styles:
        small = styles.add_style("Small note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        small = styles["Small note"]
    small.font.name = "Arial"
    small.font.size = Pt(8.5)
    small.font.color.rgb = RGBColor.from_string(MUTED)
    small.paragraph_format.space_after = Pt(3)


def set_headers_and_footers(doc: Document) -> None:
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("ValOra  |  Manuale operativo")
        r.font.name = "Arial"
        r.font.size = Pt(8.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(TEAL)

        footer = section.footer
        table = footer.add_table(rows=1, cols=2, width=Cm(17.2))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        left, right = table.rows[0].cells
        set_cell_width(left, 13.5)
        set_cell_width(right, 3.7)
        p1 = left.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run("Edizione 1.2 · 19 luglio 2026 · uso professionale")
        r1.font.name = "Arial"
        r1.font.size = Pt(7.8)
        r1.font.color.rgb = RGBColor.from_string(MUTED)
        p2 = right.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run("Pagina ")
        r2.font.name = "Arial"
        r2.font.size = Pt(7.8)
        r2.font.color.rgb = RGBColor.from_string(MUTED)
        add_page_field(p2)


def add_para(doc: Document, text: str = "", *, bold_prefix: str | None = None, style=None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def _restart_numbering(doc: Document) -> int:
    """Crea una nuova istanza della lista numerata, forzando la ripartenza da 1."""
    numbering = doc.part.numbering_part.element
    style = doc.styles["List Number"]._element
    style_num = style.find(".//" + qn("w:numId"))
    base_num_id = style_num.get(qn("w:val")) if style_num is not None else "5"
    base_num = numbering.find(f"{qn('w:num')}[@{qn('w:numId')}='{base_num_id}']")
    abstract_id = "0"
    if base_num is not None:
        abstract = base_num.find(qn("w:abstractNumId"))
        if abstract is not None:
            abstract_id = abstract.get(qn("w:val"))
    existing = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    new_id = max(existing, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return new_id


def add_bullets(doc: Document, items: list[str], *, numbered=False) -> None:
    num_id = _restart_numbering(doc) if numbered else None
    for item in items:
        p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
        if num_id is not None:
            p_pr = p._p.get_or_add_pPr()
            num_pr = p_pr.get_or_add_numPr()
            ilvl = OxmlElement("w:ilvl")
            ilvl.set(qn("w:val"), "0")
            direct_num = OxmlElement("w:numId")
            direct_num.set(qn("w:val"), str(num_id))
            num_pr.append(ilvl)
            num_pr.append(direct_num)
        p.add_run(item)


def add_callout(doc: Document, title: str, body: str, *, kind="info") -> None:
    palette = {
        "info": (PALE_BLUE, NAVY),
        "method": (MINT, TEAL),
        "warning": (PALE_GOLD, "8A5A00"),
        "limit": (PALE_RED, "9B2C2C"),
    }
    fill, color = palette[kind]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    keep_rows_together(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_width(cell, 17.0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=130, start=160, bottom=130, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(color)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.1
    p2.runs[0].font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_closing_panel(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(22)
    spacer.paragraph_format.space_after = Pt(0)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    keep_rows_together(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_width(cell, 17.0)
    set_cell_shading(cell, NAVY)
    set_cell_margins(cell, top=300, start=330, bottom=300, end=330)

    eyebrow = cell.paragraphs[0]
    eyebrow.paragraph_format.space_after = Pt(8)
    run = eyebrow.add_run("VALORA · MANUALE OPERATIVO")
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(MINT)

    statement = cell.add_paragraph()
    statement.paragraph_format.space_after = Pt(8)
    run = statement.add_run("Fonti visibili. Controlli dichiarati. Decisioni tracciabili.")
    run.bold = True
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor.from_string(WHITE)

    body = cell.add_paragraph()
    body.paragraph_format.space_after = Pt(10)
    body.paragraph_format.line_spacing = 1.12
    run = body.add_run("La tecnologia assiste il professionista; il documento originale, le regole applicate e la decisione finale restano sempre distinguibili.")
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string("D8E6ED")

    edition = cell.add_paragraph()
    edition.paragraph_format.space_after = Pt(0)
    run = edition.add_run("EDIZIONE 1.2  ·  19 LUGLIO 2026  ·  USO PROFESSIONALE")
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("AFC3CF")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = None
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        if widths:
            set_cell_width(cell, widths[i])
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        r.bold = True
        r.font.size = Pt(8.7)
        r.font.color.rgb = RGBColor.from_string(WHITE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            if widths:
                set_cell_width(cell, widths[i])
            if ridx % 2:
                set_cell_shading(cell, LIGHT)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            r.font.size = Pt(8.6)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        keep_rows_together(table.rows[-1])
    if widths:
        set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def _framed_screenshot(path: Path) -> Path:
    """Crea una copia PNG con angoli, filetto e ombra, lasciando intatto l'asset."""
    FRAMED_ASSETS.mkdir(parents=True, exist_ok=True)
    out = FRAMED_ASSETS / path.name
    if out.exists() and out.stat().st_mtime >= path.stat().st_mtime:
        return out
    with Image.open(path) as source:
        image = source.convert("RGBA")
    w, h = image.size
    radius = max(12, round(min(w, h) * 0.025))
    border = 2
    padding = 18
    shadow_offset = 6
    mask = Image.new("L", (w, h), 0)
    ImageDraw.Draw(mask).rounded_rectangle((0, 0, w - 1, h - 1), radius=radius, fill=255)
    image.putalpha(mask)

    canvas_size = (w + 2 * padding + 2 * border, h + 2 * padding + 2 * border)
    result = Image.new("RGBA", canvas_size, (255, 255, 255, 0))
    shadow_mask = Image.new("L", canvas_size, 0)
    sm = ImageDraw.Draw(shadow_mask)
    x0, y0 = padding + shadow_offset, padding + shadow_offset
    sm.rounded_rectangle((x0, y0, x0 + w + 2 * border, y0 + h + 2 * border), radius=radius + border, fill=88)
    shadow_mask = shadow_mask.filter(ImageFilter.GaussianBlur(9))
    shadow = Image.new("RGBA", canvas_size, (22, 40, 57, 0))
    shadow.putalpha(shadow_mask)
    result.alpha_composite(shadow)

    frame = Image.new("RGBA", (w + 2 * border, h + 2 * border), (216, 225, 232, 255))
    frame_mask = Image.new("L", frame.size, 0)
    ImageDraw.Draw(frame_mask).rounded_rectangle((0, 0, frame.width - 1, frame.height - 1), radius=radius + border, fill=255)
    frame.putalpha(frame_mask)
    result.alpha_composite(frame, (padding, padding))
    result.alpha_composite(image, (padding + border, padding + border))
    result.save(out, "PNG")
    return out


def add_screenshot(doc: Document, filename: str, caption: str, layout: str) -> None:
    path = ASSETS / filename
    if not path.exists():
        raise FileNotFoundError(path)
    if layout not in {"desktop", "narrow"}:
        raise ValueError(f"Layout screenshot non valido: {layout}")
    width_cm = SCREENSHOT_FULL_CM if layout == "desktop" else SCREENSHOT_NARROW_CM
    framed = _framed_screenshot(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(framed), width=Cm(width_cm))
    cp = doc.add_paragraph(caption, style="Caption")
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_flow(doc: Document, steps: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=len(steps))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    width = 17.0 / len(steps)
    for i, (number, label) in enumerate(steps):
        cell = table.cell(0, i)
        set_cell_width(cell, width)
        set_cell_shading(cell, MINT if i % 2 == 0 else PALE_BLUE)
        set_cell_margins(cell, top=120, start=80, bottom=120, end=80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(number)
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor.from_string(TEAL)
        p2 = cell.add_paragraph(label)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        p2.runs[0].font.size = Pt(8.4)
        p2.runs[0].font.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run().add_picture(str(ROOT / "public" / "logo.png"), width=Cm(5.8))
    p.paragraph_format.space_after = Pt(30)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("VALORA")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(TEAL)

    t = doc.add_paragraph(style="Title")
    t.add_run("Manuale operativo")
    st = doc.add_paragraph(style="Subtitle")
    st.add_run("Dai documenti grezzi a dati controllati, conteggi verificabili e prospetti professionali")

    line = doc.add_table(rows=1, cols=1)
    line.alignment = WD_TABLE_ALIGNMENT.LEFT
    line.autofit = False
    cell = line.cell(0, 0)
    set_cell_width(cell, 7.1)
    set_cell_shading(cell, TEAL)
    set_cell_margins(cell, top=80, start=130, bottom=80, end=130)
    band_p = cell.paragraphs[0]
    band_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    band_p.paragraph_format.space_after = Pt(0)
    band_r = band_p.add_run("EDIZIONE 1.2 · 19 LUGLIO 2026")
    band_r.bold = True
    band_r.font.name = "Arial"
    band_r.font.size = Pt(9)
    band_r.font.color.rgb = RGBColor.from_string(WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(22)

    add_para(doc, "Guida completa per studi legali, consulenti tecnici e operatori autorizzati.")
    add_para(doc, "Allineata al rilascio applicativo del 19/07 (deploy ee9ec2c)", style="Small note")
    add_para(doc, "Copre: Dynamic Island, Incidenza, Turni & Riposi, acquisizione da PC e smartphone, controlli deterministici, archivio, calcoli, esportazioni e risoluzione dei problemi.", style="Small note")

    doc.add_paragraph().paragraph_format.space_after = Pt(70)
    add_callout(
        doc,
        "Scopo del documento",
        "Questo manuale descrive il funzionamento operativo e il metodo di controllo di ValOra. Non sostituisce la valutazione giuridica, contabile o peritale del professionista, che rimane responsabile delle ipotesi adottate e del risultato finale.",
        kind="method",
    )


def add_contents(doc: Document) -> None:
    doc.add_heading("Come usare questo manuale", level=1)
    add_para(doc, "La guida segue il lavoro reale: dalla preparazione dei documenti fino alla consegna del fascicolo. La sezione 5 presenta la Dynamic Island; le sezioni 7–11 spiegano in dettaglio acquisizione, estrazione e controlli; le sezioni successive coprono revisione, calcoli, archivio ed export.")
    add_table(
        doc,
        ["Parte", "Contenuto", "Sezioni"],
        [
            ["A · Orientamento", "Ambiti, ruoli, flusso generale e Dynamic Island", "1–5"],
            ["B · Acquisizione", "Creazione lavoratore, caricamento PC/cartella e scanner mobile", "6–8"],
            ["C · Dato e prova", "Estrazione AI, parser deterministico, anomalie e verifica", "9–12"],
            ["D · Lavorazione", "Revisione mensile, archivio, incidenza, TFR e riposi", "13–17"],
            ["E · Consegna", "Report, demo, limiti, assistenza, checklist e glossario", "18–22"],
        ],
        [3.0, 10.7, 3.3],
    )
    doc.add_heading("Legenda", level=2)
    legend = add_table(
        doc,
        ["Segnale", "Significato operativo"],
        [
            ["Metodo", "Regola strutturale del flusso ValOra."],
            ["Attenzione", "Passaggio che richiede una decisione o una verifica dell’operatore."],
            ["Limite attuale", "Funzione non ancora completa o dato non aggiornato in tempo reale."],
        ],
        [4.0, 13.0],
    )
    chip_palette = [("Metodo", MINT, TEAL), ("Attenzione", PALE_GOLD, "8A5A00"), ("Limite attuale", PALE_RED, "9B2C2C")]
    for row, (text, fill, color) in zip(legend.rows[1:], chip_palette):
        cell = row.cells[0]
        set_cell_shading(cell, WHITE)
        p = cell.paragraphs[0]
        p.clear()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(f"  {text}  ")
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(color)
        set_run_shading(run, fill)
    add_callout(doc, "Principio guida", "ValOra accelera e rende ripetibile il lavoro documentale; l’operatore mantiene il controllo sulle correzioni, sulle ipotesi di calcolo e sulla validazione finale.", kind="method")


def build_manual() -> Document:
    doc = Document()
    set_doc_defaults(doc)
    add_cover(doc)
    page_break(doc)
    add_contents(doc)

    page_break(doc)
    doc.add_heading("1. Che cos’è ValOra", level=1)
    add_para(doc, "ValOra è il motore LegalTech proprietario che organizza cedolini, turni e dati retributivi in pratiche verificabili. È progettato per supportare il lavoro professionale: conserva il collegamento fra documento originale, dato estratto, correzioni e risultato di calcolo.")
    add_bullets(doc, [
        "Incidenza: analisi di ferie, permessi, festività, voci fisse e variabili, differenze e TFR.",
        "Turni & Riposi: ricostruzione cronologica e confronto fra fonte documentale e classificazione elaborata.",
        "Archivio: conservazione e consultazione dei documenti associati ai singoli mesi.",
        "Report: prospetti di lavoro, riepiloghi ed esportazioni per il fascicolo professionale.",
    ])
    add_callout(doc, "Posizionamento", "ValOra non è un semplice lettore di buste paga e non sostituisce il professionista. È l’infrastruttura metodologica che rende il processo più veloce, tracciabile e controllabile.", kind="info")

    doc.add_heading("2. Aree di lavoro e stati", level=1)
    add_table(doc, ["Area", "Quando usarla", "Risultato atteso"], [
        ["Incidenza", "Quando occorre ricostruire la retribuzione utile e verificare ferie, permessi, festività e TFR.", "Griglia mensile, differenze annuali, prospetti ed export."],
        ["Turni & Riposi", "Quando il problema dipende dalla successione dei turni e dal riposo tra giornate lavorate.", "Cronologia confrontabile e prospetti dei riposi contestabili."],
        ["Indennità", "Area di studio per ulteriori istituti economici.", "Anteprima funzionale; non usarla come modulo produttivo definitivo."],
    ], [3.4, 7.2, 6.4])
    add_table(doc, ["Stato pratica (etichetta nell’app)", "Significato"], [
        ["Da Analizzare", "Pratica aperta: dati o documenti ancora incompleti, lavorazione da avviare o completare."],
        ["Conteggi", "Dati revisionati; la pratica è in fase di calcolo e produzione dei prospetti."],
        ["Buste Paga Mancanti", "La pratica è stata inviata/istruita ma mancano documenti per completare l’arco temporale."],
        ["Conclusa", "Lavorazione completata; fascicolo definito, in attesa dell’esito."],
        ["Pagata", "Pratica liquidata: chiusa a tutti gli effetti (per il consultatore è anche la condizione per il download)."],
    ], [5.0, 12.0])
    add_callout(doc, "Attenzione", "Usare sempre le etichette sopra, che sono quelle visibili nella timeline «Stato vertenza» e nei raccoglitori della dashboard: stati con altri nomi non esistono nell’interfaccia.", kind="warning")
    add_screenshot(doc, "02-scelta-area.png", "Schermata di scelta dell’ambito operativo nella demo ValOra.", "desktop")

    doc.add_heading("3. Accesso, organizzazione e ruoli", level=1)
    add_para(doc, "Ogni attività si svolge nell’organizzazione selezionata. Ogni operatore accede soltanto ai dati del proprio account, con autorizzazioni controllate anche lato server; la scelta dell’organizzazione filtra le pratiche mostrate e le pratiche precedenti alla gestione multi-organizzazione possono comparire in ogni organizzazione.")
    add_table(doc, ["Ruolo", "Può fare", "Non può fare"], [
        ["Owner / operatore", "Creare e modificare pratiche, acquisire documenti, correggere dati, avviare verifiche ed esportare.", "Accedere ai dati di organizzazioni per cui non è autorizzato."],
        ["Viewer", "Consultare pratiche, prospetti e documenti consentiti.", "Modificare dati, avviare importazioni o approvare correzioni."],
    ], [3.3, 7.1, 6.6])
    add_bullets(doc, [
        "Prima di lavorare, controllare organizzazione e area selezionate.",
        "Non condividere link temporanei ai documenti oltre il destinatario autorizzato.",
        "Usare la demo per formazione o presentazioni, non per archiviare dati reali.",
    ])

    doc.add_heading("4. Il flusso completo", level=1)
    add_flow(doc, [("1", "Prepara"), ("2", "Acquisisci"), ("3", "Estrai"), ("4", "Controlla"), ("5", "Calcola"), ("6", "Consegna")])
    add_table(doc, ["Fase", "Cosa accade", "Controllo umano"], [
        ["Prepara", "Si crea il lavoratore e si ordinano i file disponibili.", "Identità, azienda, arco temporale e completezza."],
        ["Acquisisci", "PDF o immagini entrano da PC, cartella o smartphone.", "Mese/anno proposto e qualità del documento."],
        ["Estrai", "Il profilo AI aziendale restituisce dati strutturati.", "Coerenza visiva con il cedolino."],
        ["Controlla", "Regole deterministiche e verifiche locali segnalano discrepanze.", "Accettazione, rifiuto o correzione."],
        ["Calcola", "Le tabelle applicano il metodo della pratica.", "Ipotesi, inclusioni, cap e dati mancanti."],
        ["Consegna", "Si producono prospetti, PDF, DOCX, ZIP e, per Turni & Riposi, Excel.", "Revisione finale e conservazione del fascicolo."],
    ], [2.7, 7.3, 7.0])

    doc.add_heading("5. Dynamic Island: il centro operativo di ValOra", level=1)
    add_para(doc, "La Dynamic Island è una superficie persistente e adattiva che accompagna il lavoro senza coprire la pratica. Riunisce attività in background, notifiche, strumenti e azioni contestuali in un unico punto sempre visibile. Durante le acquisizioni si comporta come una Live Activity: il documento non scompare dietro uno spinner, ma resta rappresentato fino all’esito.")
    add_callout(doc, "Perché è un’innovazione", "ValOra porta nel software documentale e LegalTech un paradigma tipico delle Live Activity: avanzamento, continuità PC–smartphone e strumenti contestuali restano comprensibili e richiamabili mentre l’operatore continua a lavorare.", kind="method")
    add_screenshot(doc, "17-dynamic-island-menu.png", "Dynamic Island espansa nella scheda lavoratore: un centro di comando sovrapposto al contesto, senza abbandonare la pratica.", "desktop")
    add_bullets(doc, [
        "Stato compatto: accesso immediato a menu, calcolatrice e assistenti.",
        "Stato contestuale: scorciatoie diverse nella scheda, durante lo scroll e nel report.",
        "Stato informativo: ticker del netto, notifiche, esiti e scheda sintetica del lavoratore.",
        "Stato operativo: Live Activity per file singolo, batch, cartella e scanner mobile.",
    ])

    page_break(doc)
    doc.add_heading("5.1 Live Activity e continuità del lavoro", level=2)
    add_flow(doc, [("1", "Avvio"), ("2", "Documento corrente"), ("3", "Avanzamento"), ("4", "Esito"), ("5", "Notifica")])
    add_table(doc, ["Modalità", "Cosa mostra l’Island", "Continuità operativa"], [
        ["File singolo", "Stato della scansione, documento corrente e completamento.", "La pratica resta visibile; al termine compare l’esito."],
        ["Batch", "Elementi elaborati e in attesa, file corrente, ETA stimata e progresso.", "La Live Activity può essere aperta o minimizzata."],
        ["Cartella", "Conteggio, percentuale, file corrente e tema dedicato.", "Il lavoro continua in background senza bloccare la scheda."],
        ["Mobile QR", "Stato di ascolto, fascicoli ricevuti, ultimi eventi ed esito.", "Il controllo passa dal modale QR all’Island sul computer."],
    ], [3.2, 7.0, 6.8])
    add_screenshot(doc, "13-dynamic-island-viewport.png", "La pill dell’Island mostra il totale netto nella pratica e cambia forma quando interviene un’attività o uno strumento.", "desktop")
    doc.add_heading("Minimizzare senza interrompere", level=2)
    add_para(doc, "Durante un caricamento l’operatore può ridurre la Live Activity a una pill satellite. Il job prosegue e la pill conserva progresso, anello visivo, documento corrente ed ETA per i caricamenti non mobile. Un clic ripristina la vista completa; al completamento l’esito torna automaticamente in primo piano.")
    add_bullets(doc, [
        "Un upload attivo non minimizzato ha priorità sulle notifiche secondarie.",
        "Se il progresso resta fermo per oltre 20 secondi, compare l’avviso Verifica…: è un segnale, non un retry automatico.",
        "Un nuovo upload annulla il timer dell’esito precedente, evitando messaggi riferiti al lavoro sbagliato.",
        "La calcolatrice mantiene gli ultimi dieci calcoli; menu e quick action evitano continui cambi di pagina.",
    ])
    add_callout(doc, "Limiti attuali", "L’Island segue un job alla volta, può minimizzare ma non annullare l’upload e non mostra l’ETA mobile. I testi di avanzamento descrivono la pipeline, non i controlli deterministici.", kind="limit")

    doc.add_heading("6. Creare la scheda del lavoratore", level=1)
    add_bullets(doc, [
        "Aprire l’area corretta e scegliere Nuova pratica / Nuovo lavoratore.",
        "Inserire i dati identificativi strettamente necessari e selezionare l’azienda o il profilo documentale.",
        "Definire il periodo da analizzare. Un periodo realistico evita mesi vuoti e controlli fuorvianti.",
        "Salvare, quindi aprire la scheda per acquisire documenti o compilare manualmente i mesi mancanti.",
    ], numbered=True)
    add_callout(doc, "Prima dell’importazione", "Uniformare i nomi dei file quando possibile, rimuovere doppioni evidenti e verificare che ogni pagina sia leggibile. ValOra può gestire file eterogenei, ma una fonte ordinata riduce i casi da verificare.", kind="warning")
    add_screenshot(doc, "04-scheda-lavoratore.png", "Scheda lavoratore: punto centrale per dati mensili, documenti e azioni di verifica.", "narrow")

    page_break(doc)
    doc.add_heading("7. Acquisire le buste paga dal computer", level=1)
    add_para(doc, "ValOra accetta PDF e immagini JPEG, PNG, WebP e HEIC. I file nascosti vengono ignorati. L’importazione ordinaria usa lo stesso percorso di estrazione AI sia per PDF testuali sia per scansioni: il parser deterministico non è il motore di importazione.")
    add_table(doc, ["Modalità", "Uso consigliato", "Comportamento"], [
        ["Singolo file", "Correzione puntuale o un mese isolato.", "Si seleziona il documento e si controlla subito il periodo proposto."],
        ["Più file", "Piccoli gruppi già ordinati.", "I documenti vengono elaborati in parallelo; i risultati arrivano progressivamente."],
        ["Intera cartella", "Serie storiche e fascicoli organizzati per anno/mese.", "ValOra percorre i file compatibili e usa anche il nome della cartella immediatamente superiore per interpretare il periodo."],
    ], [3.2, 6.2, 7.6])
    doc.add_heading("Come viene individuato il mese", level=2)
    add_bullets(doc, [
        "Prima priorità: mese e anno riconoscibili nel nome del file.",
        "Seconda priorità: indicazioni presenti nella cartella immediatamente superiore.",
        "Confronto finale: il periodo proposto viene confrontato con quello estratto dal documento.",
        "Se le indicazioni non coincidono, ValOra non sostituisce silenziosamente il dato: marca il caso come Mese da verificare.",
    ])
    add_callout(doc, "Robustezza del batch", "L’elaborazione può ritentare automaticamente fino a tre volte i file che falliscono. Un errore di archiviazione non annulla i dati già estratti; il documento deve però essere reinserito nell’archivio se il caricamento non è riuscito.", kind="info")

    page_break(doc)
    doc.add_heading("8. Acquisire da smartphone con QR", level=1)
    add_para(doc, "Lo scanner mobile consente di fotografare una busta paga o scegliere un PDF/immagine dal telefono senza trasferimenti manuali. Il computer crea una sessione privata e temporanea; il QR collega lo smartphone alla pratica in corso. Quando il telefono passa all’elaborazione, il grande modale QR può chiudersi e il controllo continua nella Dynamic Island sul computer.")
    add_bullets(doc, [
        "Dal computer, aprire la scheda e scegliere l’acquisizione mobile.",
        "Inquadrare il QR con lo smartphone e aprire la sessione proposta.",
        "Fotografare tutte le pagine nell’ordine corretto oppure scegliere un documento esistente.",
        "Controllare nitidezza, orientamento e margini; quindi inviare.",
        "Attendere sul computer il risultato. I file falliti restano disponibili per un nuovo tentativo.",
    ], numbered=True)
    add_screenshot(doc, "05-scanner-mobile-qr.png", "Sessione mobile privata aperta tramite QR. La demo usa dati e flussi dimostrativi.", "narrow")
    add_table(doc, ["Buona acquisizione", "Da evitare"], [
        ["Pagina intera, luce uniforme, testo nitido, telefono parallelo al foglio.", "Ombre, riflessi, tagli, movimento, pieghe sul testo o più cedolini nella stessa foto."],
        ["Una sequenza completa per documento multipagina.", "Pagine invertite o documenti diversi uniti nello stesso invio."],
    ], [8.5, 8.5])
    add_callout(doc, "Durata della sessione", "La sessione mobile scade e il computer interrompe l’attesa dopo inattività. Se il collegamento è scaduto, generare un nuovo QR dalla scheda.", kind="warning")

    page_break(doc)
    doc.add_heading("9. Come vengono estratti i dati", level=1)
    add_para(doc, "Dopo l’acquisizione, ValOra invia il contenuto del documento al motore AI con un profilo specializzato per azienda. Il profilo definisce terminologia, struttura attesa e campi da restituire. La risposta viene normalizzata in JSON e associata al mese della pratica.")
    add_flow(doc, [("A", "File originale"), ("B", "Profilo azienda"), ("C", "Estrazione AI"), ("D", "JSON strutturato"), ("E", "Griglia mensile")])
    add_table(doc, ["Profilo", "Specializzazione"], [
        ["RFI / Trenitalia", "Voci ferroviarie, presenze, ferie/permessi e struttura dei relativi cedolini."],
        ["FSE", "Formati e voci Ferrovie del Sud Est, con supporto alle verifiche locali per i documenti compatibili."],
        ["Mercitalia", "Formati ADP e dati utili alle quadrature previste dal verificatore."],
        ["Elior", "Cedolini spesso acquisiti come scansione; al risultato OCR si applicano controlli aritmetici dedicati."],
        ["Clean Service / personalizzati", "Prompt e mappature adattati alla fonte disponibile."],
        ["Generico", "Fallback per documenti non riconosciuti; richiede una revisione più attenta."],
    ], [4.4, 12.6])
    add_callout(doc, "Distinzione essenziale", "L’AI legge e struttura il documento. Il parser deterministico è un secondo livello, separato e disponibile soltanto per documenti/profili supportati. Non bisogna descriverlo come il metodo con cui ValOra importa normalmente i PDF testuali.", kind="method")
    doc.add_heading("Dati tipicamente acquisiti", level=2)
    add_bullets(doc, [
        "periodo del cedolino e dati anagrafici utili all’associazione;",
        "giorni o ore di ferie, permessi, festività e altre assenze;",
        "voci fisse, voci variabili, trattenute e importi già corrisposti;",
        "presenze e quantità necessarie ai controlli previsti dal profilo;",
        "totali utili alla quadratura, quando esposti dal documento.",
    ])

    page_break(doc)
    doc.add_heading("10. Il controllo deterministico", level=1)
    add_para(doc, "Un controllo deterministico applica regole dichiarate e ripetibili: a parità di documento e configurazione produce lo stesso confronto. Non interpreta liberamente il testo e non sostituisce il giudizio dell’operatore.")
    add_table(doc, ["Livello", "Input", "Cosa controlla", "Esito"], [
        ["Estrazione AI", "PDF o immagine", "Individua e struttura i dati secondo il profilo aziendale.", "Valori proposti nella pratica."],
        ["Validazione deterministica in import", "Dati estratti / OCR", "Coerenze aritmetiche e logiche previste dal profilo.", "Avvisi, anomalie o correzioni sicure limitate."],
        ["Verifica accuratezza dal disco", "PDF testuale scelto dall’utente", "Confronto locale tra colonne/valori del PDF e dati salvati.", "Differenze da accettare o respingere."],
        ["Seconda verifica AI", "PDF archiviato + dati JSON", "Rilettura assistita del dato già acquisito.", "Suggerimenti da approvare; non è il parser deterministico."],
    ], [3.5, 3.5, 6.1, 3.9])
    doc.add_heading("Controlli durante l’importazione", level=2)
    add_bullets(doc, [
        "RFI/Trenitalia: coerenza delle presenze, somme plausibili, possibili slittamenti di colonna e valori oltre i limiti del calendario. Solo gli slittamenti considerati sicuri possono essere corretti automaticamente; gli altri casi vengono segnalati.",
        "Elior: prodotto quantità × tariffa rispetto alla competenza, somma delle competenze rispetto ai totali, rapporto competenze − trattenute ≈ netto, giorni e tariffe plausibili. Il validatore segnala: non modifica automaticamente i dati.",
        "I controlli dipendono dai campi effettivamente disponibili. Se righe o totali mancano nel risultato di origine, alcune verifiche non possono essere eseguite.",
    ])
    add_callout(doc, "Prova e tracciabilità", "La forza del controllo deterministico non è promettere l’assenza assoluta di errori, ma rendere esplicite le regole, mostrare le discrepanze e lasciare all’operatore la decisione finale.", kind="method")

    page_break(doc)
    doc.add_heading("11. Verifica accuratezza dal disco", level=1)
    add_para(doc, "Questa funzione è la prova locale e indipendente dall’AI per i profili supportati. Si avvia dalla scheda del lavoratore scegliendo Azioni → Verifica accuratezza (dal disco), quindi selezionando i PDF originali.")
    add_bullets(doc, [
        "Il PDF viene letto localmente nel browser con PDF.js: non viene inviato al servizio AI per questa verifica.",
        "Il parser aziendale confronta i valori stampati nelle colonne con quelli già salvati nella pratica.",
        "Le differenze oltre la tolleranza tecnica vengono presentate come proposte; l’utente decide quali applicare.",
        "I dati non vengono modificati fino alla conferma esplicita dell’operatore.",
    ])
    add_screenshot(doc, "07-verifica-accuratezza.png", "Pannello dimostrativo della verifica di accuratezza sui documenti selezionati dal disco.", "narrow")
    add_table(doc, ["Profilo supportato", "Ambito della verifica"], [
        ["RFI / Trenitalia", "PDF testuali: confronto delle colonne e controlli di presenza; il totale stampato non è disponibile in ogni formato."],
        ["FSE", "PDF testuali compatibili, in particolare dai formati successivi a luglio 2017; quadratura quando i totali necessari sono presenti."],
        ["Mercitalia", "PDF ADP compatibili, con controlli di quadratura previsti dal profilo."],
        ["Elior / scansioni", "Non usa questo parser testuale: si applicano i controlli aritmetici deterministici sull’output OCR durante l’importazione."],
    ], [5.0, 12.0])
    page_break(doc)
    doc.add_heading("Quando la verifica si ferma", level=2)
    add_bullets(doc, [
        "PDF scansionato senza testo, cifrato o con struttura non supportata;",
        "tredicesima o quattordicesima che non deve essere confusa con un mese ordinario;",
        "periodo del file non coerente con la pratica o mese assente;",
        "quadratura non superata, valori in conflitto o giorni ambigui;",
        "mese non ancora presente nella griglia: la verifica controlla, ma non crea automaticamente il periodo mancante.",
    ])
    add_callout(doc, "Tolleranza", "I confronti numerici usano una tolleranza tecnica molto stretta (0,005) per assorbire le sole differenze di arrotondamento. Una differenza sostanziale resta visibile.", kind="info")

    doc.add_heading("12. Anomalie, scarti e decisioni dell’operatore", level=1)
    add_table(doc, ["Segnale", "Causa probabile", "Azione consigliata"], [
        ["Mese da verificare", "Nome file/cartella e periodo letto nel cedolino non coincidono.", "Aprire il PDF, identificare il periodo corretto e confermare o spostare i dati."],
        ["Documento non leggibile", "Scansione sfocata, tagliata, protetta o formato non supportato.", "Ripetere la scansione o ottenere il PDF originale."],
        ["Valore fuori scala", "OCR errato, colonna slittata o dato realmente anomalo.", "Confrontare con il documento e correggere solo dopo verifica."],
        ["Quadratura fallita", "Totali e righe non tornano entro la tolleranza.", "Controllare trattenute, segni, quantità e pagine mancanti."],
        ["Archivio assente", "Estrazione riuscita ma caricamento dell’originale fallito.", "Reinserire manualmente il documento nell’archivio del mese."],
        ["Doppione", "Lo stesso mese o lo stesso documento è stato acquisito più volte.", "Confrontare versioni e mantenere quella corretta; l’import non sostituisce automaticamente l’originale esistente."],
    ], [4.0, 6.2, 6.8])
    add_callout(doc, "Regola di sicurezza", "Un avviso non dimostra da solo che il cedolino sia errato: indica che la catena documentale non è sufficientemente coerente per procedere senza controllo umano.", kind="warning")

    doc.add_heading("13. Revisione mensile e viewer affiancato", level=1)
    add_para(doc, "La griglia mensile è l’ambiente di revisione. Le celle possono essere corrette singolarmente, incollate da fonti esterne compatibili, compilate per intervallo o annullate. Il viewer consente di tenere il documento accanto ai dati.")
    add_bullets(doc, [
        "Aprire un mese e visualizzare il cedolino originale.",
        "Controllare periodo, presenza delle voci e segno degli importi.",
        "Verificare ferie, permessi e festività separatamente dalle voci economiche.",
        "Applicare soltanto i suggerimenti confermati dal documento.",
        "Annotare eccezioni metodologiche nella pratica prima del calcolo finale.",
    ], numbered=True)
    add_callout(doc, "Uso mobile", "Sul telefono la griglia mensile è in sola consultazione: elenco dei mesi con dettaglio voci, senza modifica dei dati. Dal rilascio del 19 luglio 2026 su smartphone si consultano anche Riepilogo Annuale, Analisi Voci, archivio, visore a pagina intera e report; le correzioni si fanno da computer o tablet.", kind="info")

    doc.add_heading("14. Archivio documentale", level=1)
    add_para(doc, "Ogni documento può essere associato al mese e conservato con i metadati della pratica. L’archivio permette apertura, download e preparazione di pacchetti ZIP. I collegamenti di visualizzazione sono firmati e temporanei.")
    add_screenshot(doc, "08-archivio.png", "Archivio documentale mensile nella demo: originali, stato e azioni disponibili.", "narrow")
    add_table(doc, ["Operazione", "Effetto"], [
        ["Archivia durante l’import", "Salva l’originale in background e collega il risultato estratto al mese."],
        ["Archivia manualmente", "Conserva il documento, ma non compila da sola la griglia mensile."],
        ["Rianalizza", "Aggiorna i dati estratti conservati nell’archivio; non riscrive automaticamente i valori già presenti nella griglia."],
        ["Scarica ZIP", "Raccoglie i documenti disponibili per la consegna o la conservazione esterna."],
    ], [5.2, 11.8])
    add_callout(doc, "Attenzione ai duplicati", "Se esiste già un documento per lo stesso mese, una nuova scansione non sostituisce automaticamente l’originale. Confrontare i file e scegliere consapevolmente quale conservare.", kind="warning")

    doc.add_heading("15. Incidenza: logica di calcolo", level=1)
    add_para(doc, "Il modulo Incidenza distingue la retribuzione fissa dalla componente variabile. Le voci fisse costituiscono la base di confronto; il credito deriva dalla media giornaliera delle voci variabili secondo le impostazioni della pratica.")
    add_table(doc, ["Passaggio", "Formula operativa semplificata"], [
        ["Media variabile giornaliera", "Somma delle voci variabili considerate ÷ giorni utili del periodo."],
        ["Media applicabile", "Di regola si usa la media dell’anno precedente; in assenza, ValOra può usare l’anno corrente e segnala il fallback."],
        ["Giornate utili", "Ferie/istituti ammessi, normalmente entro 28 giorni; fino a 32 includendo le ex festività quando previsto."],
        ["Importo teorico", "Media giornaliera × giornate utili."],
        ["Differenza netta", "Importo teorico − importi già corrisposti − eventuali ticket considerati."],
    ], [5.0, 12.0])
    add_callout(doc, "Ipotesi da dichiarare", "L’inclusione di permessi, festività, ticket e singole voci variabili dipende dal quesito e dalla metodologia adottata. Le impostazioni devono essere riportate nel prospetto o nella relazione.", kind="warning")

    doc.add_heading("16. Incidenza, TFR e rivalutazione", level=1)
    doc.add_heading("Indicatori di incidenza", level=2)
    add_para(doc, "Per ciascun mese ValOra può rapportare la componente variabile alla retribuzione complessiva utile. L’incidenza annuale è ottenuta aggregando i mesi disponibili; i periodi incompleti devono essere segnalati, perché possono alterare la media.")
    doc.add_heading("TFR sulle differenze", level=2)
    add_bullets(doc, [
        "Si individua la differenza retributiva progressiva maturata nell’anno.",
        "La quota TFR viene calcolata secondo il divisore 13,5 e la contribuzione dello 0,5% prevista dal modello.",
        "Le quote pregresse possono essere rivalutate usando il coefficiente previsto: componente fissa dell’1,5% più il 75% della variazione dell’indice applicabile.",
        "I risultati devono essere letti insieme alle date di maturazione e alle ipotesi della pratica.",
    ])
    add_callout(doc, "Limite attuale TFR", "La tabella interna dei coefficienti TFR è censita fino al 2024. Per periodi successivi il sistema usa un fallback tecnico dell’1,5%: prima di consegnare una pratica aggiornata occorre inserire o verificare i coefficienti ufficiali.", kind="limit")
    doc.add_heading("Indici ISTAT", level=2)
    add_para(doc, "ValOra utilizza una tabella interna degli indici censiti, aggiornata fino all’ultima pubblicazione registrata nel progetto (maggio 2026). Il controllo di rete verifica la disponibilità della fonte, ma non importa automaticamente l’ultimo indice: l’aggiornamento non è in tempo reale.")

    page_break(doc)
    doc.add_heading("17. Turni & Riposi", level=1)
    add_para(doc, "Il modulo ricostruisce la sequenza dei turni e mette a confronto il dato di fonte con la classificazione prodotta dal motore. Le due letture restano separate: il sistema non le somma automaticamente e rende visibili le divergenze.")
    add_bullets(doc, [
        "Aprire una pratica Turni & Riposi già esistente.",
        "Verificare la sequenza cronologica e la completezza dei periodi.",
        "Confrontare classificazione della fonte e classificazione calcolata.",
        "Controllare eccezioni, riposi e importi prima dell’esportazione.",
        "Produrre Excel, stampa/PDF o DOCX secondo la destinazione del fascicolo.",
    ], numbered=True)
    add_screenshot(doc, "06-turni-riposi.png", "Area Turni & Riposi nella demo ValOra.", "narrow")
    add_callout(doc, "Limite attuale", "La creazione di una nuova pratica e l’importazione automatica dei turni non sono ancora disponibili nell’interfaccia corrente. Il modulo è utilizzabile sulle pratiche già predisposte.", kind="limit")

    doc.add_heading("18. Report ed esportazioni", level=1)
    add_para(doc, "La consegna deve permettere al professionista di ricostruire il passaggio dai documenti ai risultati. ValOra offre prospetti stampabili e pacchetti esportabili; la composizione effettiva dipende dall’area e dallo stato della pratica.")
    add_table(doc, ["Formato", "Contenuto tipico", "Uso"], [
        ["PDF conteggi", "Tabelle mensili/annuali, differenze, totali e note metodologiche.", "Deposito, condivisione e revisione stabile."],
        ["PDF riepilogo", "Dati essenziali della pratica e risultato sintetico.", "Orientamento rapido del legale o del cliente."],
        ["DOCX tecnico", "Relazione modificabile e sezioni narrative.", "Revisione professionale e personalizzazione."],
        ["Excel", "Disponibile per Turni & Riposi: dati strutturati con formule verificabili.", "Controlli ulteriori e scambio con altri strumenti."],
        ["ZIP", "Originali e/o più elaborati raccolti insieme.", "Consegna del fascicolo e conservazione."],
    ], [3.3, 7.8, 5.9])
    add_screenshot(doc, "09-prospetto-ufficiale.png", "Esempio dimostrativo di prospetto ufficiale generato da ValOra.", "narrow")
    add_bullets(doc, [
        "Prima dell’export, filtrare o annotare i mesi incompleti.",
        "Verificare che il periodo del report coincida con il mandato.",
        "Controllare totali e formule sul prospetto finale, non solo nella griglia.",
        "Conservare insieme elaborato, originali e nota delle ipotesi adottate.",
    ])

    doc.add_heading("19. Modalità demo", level=1)
    add_para(doc, "La demo consente di presentare il flusso senza usare dati personali e senza dipendere dai servizi reali. Simula l’estrazione, mostra pratiche dimostrative e permette di osservare il comportamento adattivo della Dynamic Island; non sostituisce un test produttivo del motore AI, del database o dell’archiviazione.")
    add_screenshot(doc, "01-ingresso-valora.png", "Ingresso alla demo ValOra e accesso alle aree dimostrative.", "narrow")
    add_callout(doc, "Uso corretto della demo", "Usarla per formazione, presentazioni e prova dell’interfaccia. Non considerare i tempi o i risultati simulati come misura dell’accuratezza su un nuovo formato di cedolino.", kind="info")

    doc.add_heading("20. Limiti attuali e funzioni in evoluzione", level=1)
    add_table(doc, ["Ambito", "Situazione attuale", "Conseguenza operativa"], [
        ["Indennità", "Area in anteprima con classificazioni ancora oggetto di consolidamento.", "Non usarla come calcolo produttivo definitivo."],
        ["Turni & Riposi", "Nuova pratica e importazione non ancora abilitate nella UI.", "Operare su pratiche già predisposte."],
        ["Rivalutazione TFR", "Coefficienti interni completi fino al 2024; fallback dopo tale data.", "Verificare i coefficienti ufficiali prima della consegna."],
        ["ISTAT", "Tabella interna aggiornata all’ultima pubblicazione censita, non sincronizzata in tempo reale.", "Controllare l’indice per pratiche recentissime."],
        ["Rianalisi archivio", "Aggiorna i dati estratti del documento, non la griglia mensile.", "Riportare manualmente le correzioni approvate."],
        ["PDF scansionati", "La verifica locale testuale non è applicabile.", "Affidarsi all’estrazione AI + validazioni disponibili + controllo visivo."],
        ["Selettori temporali", "Alcune viste legacy terminano al 2025.", "Usare le viste aggiornate o verificare manualmente i periodi successivi."],
    ], [4.1, 6.5, 6.4])
    add_callout(doc, "Trasparenza", "Indicare nel fascicolo se un passaggio è stato verificato da parser deterministico, soltanto da controlli aritmetici o esclusivamente mediante revisione visiva. I tre livelli non sono equivalenti.", kind="method")

    doc.add_heading("21. Risoluzione dei problemi", level=1)
    add_table(doc, ["Problema", "Controlli rapidi", "Soluzione"], [
        ["Il file non parte", "Estensione, dimensione, cifratura e connessione.", "Convertire in PDF/PNG supportato, rimuovere la protezione se autorizzati e riprovare."],
        ["Il mese è sbagliato", "Nome file, cartella e periodo stampato.", "Correggere il nome o confermare manualmente il mese nella pratica."],
        ["Il QR non si collega", "Scadenza della sessione, stessa pratica ancora aperta, rete mobile.", "Chiudere la sessione e generare un nuovo QR."],
        ["Mancano voci", "Qualità immagine, pagina mancante, profilo aziendale.", "Ripetere acquisizione o integrare manualmente dopo confronto con l’originale."],
        ["Quadratura non torna", "Segni, trattenute, quantità/tariffe, totale e pagine.", "Correggere soltanto il dato documentato; mantenere una nota se la fonte è incoerente."],
        ["Il PDF non appare in archivio", "Stato del caricamento e duplicati del mese.", "Archiviare manualmente il file; verificare prima quale versione conservare."],
        ["La rianalisi non cambia la tabella", "Dato estratto dell’archivio vs griglia pratica.", "Aprire il confronto e applicare manualmente le variazioni approvate."],
    ], [4.0, 5.6, 7.4])

    doc.add_heading("22. Checklist operative", level=1)
    doc.add_heading("Prima di acquisire", level=2)
    add_bullets(doc, [
        "□ Organizzazione, area, lavoratore e azienda sono corretti.",
        "□ Il periodo della pratica corrisponde al mandato.",
        "□ I file sono leggibili, completi e privi di doppioni evidenti.",
        "□ Nomi e cartelle contengono mese/anno quando possibile.",
    ])
    doc.add_heading("Dopo l’estrazione", level=2)
    add_bullets(doc, [
        "□ Mese e anno coincidono con il documento.",
        "□ Le principali voci economiche e le presenze sono state confrontate.",
        "□ Gli avvisi deterministici sono stati risolti o motivati.",
        "□ Il PDF originale è presente nell’archivio.",
        "□ Le correzioni sono state applicate soltanto dopo conferma visiva.",
    ])
    doc.add_heading("Prima della consegna", level=2)
    add_bullets(doc, [
        "□ L’arco temporale è completo o le lacune sono dichiarate.",
        "□ Le ipotesi di incidenza, cap, ticket e istituti inclusi sono esplicite.",
        "□ Coefficienti TFR e indici recenti sono stati verificati.",
        "□ Totali del prospetto finale e dati della pratica coincidono.",
        "□ Originali, prospetti e relazione formano un fascicolo coerente.",
    ])
    add_callout(doc, "Chiusura della pratica", "Una pratica è pronta quando un altro professionista può ripercorrere le fonti, comprendere le scelte e ottenere gli stessi passaggi di controllo senza affidarsi a informazioni non documentate.", kind="method")

    doc.add_heading("Glossario essenziale", level=1)
    add_table(doc, ["Termine", "Definizione"], [
        ["Acquisizione", "Ingresso di PDF o immagini da PC, cartella o smartphone."],
        ["Estrazione AI", "Lettura assistita del documento e trasformazione in dati strutturati secondo un profilo."],
        ["OCR", "Riconoscimento del testo presente in un’immagine o scansione."],
        ["Parser deterministico", "Lettore basato su regole esplicite che produce lo stesso risultato sullo stesso input."],
        ["Validazione", "Controllo logico o aritmetico sui dati estratti; può segnalare senza correggere."],
        ["Quadratura", "Confronto tra somme delle righe e totali stampati, entro la tolleranza prevista."],
        ["Profilo aziendale", "Configurazione di estrazione e controllo adattata al formato del datore di lavoro."],
        ["Griglia mensile", "Tabella della pratica che contiene i valori utilizzati nei conteggi."],
        ["Archivio", "Raccolta degli originali e dei relativi metadati, separata dalla griglia di calcolo."],
        ["Fallback", "Regola sostitutiva usata quando manca il dato preferenziale; deve essere segnalata e verificata."],
        ["Dynamic Island", "Centro operativo persistente che mostra attività, progresso, strumenti, azioni contestuali ed esiti senza interrompere la pratica."],
        ["Live Activity", "Stato espanso dell’Island che accompagna un caricamento dal file corrente fino al risultato e può essere minimizzato."],
    ], [4.6, 12.4])

    doc.add_heading("Appendice A · Matrice dei controlli", level=1)
    add_table(doc, ["Fonte", "Estrazione ordinaria", "Controllo deterministico", "Intervento umano"], [
        ["PDF testuale supportato", "AI con profilo aziendale.", "Verifica locale dal disco; quadratura dove disponibile.", "Conferma delle differenze e delle ipotesi."],
        ["PDF testuale generico", "AI con profilo generico o personalizzato.", "Non disponibile se il formato non ha un parser dedicato.", "Revisione visiva rafforzata."],
        ["Scansione Elior", "AI/OCR con profilo Elior.", "Controlli aritmetici sull’output OCR.", "Confronto visivo e correzione delle anomalie."],
        ["Altra scansione", "AI/OCR con profilo disponibile.", "Solo le validazioni previste dal profilo; niente parser testuale locale.", "Revisione visiva completa."],
        ["Inserimento manuale", "Nessuna estrazione.", "Controlli della griglia e del calcolo, se applicabili.", "Responsabilità integrale dell’inserimento."],
    ], [4.0, 4.0, 5.0, 4.0])
    add_callout(doc, "Formula corretta da usare nelle presentazioni", "ValOra acquisisce PDF e immagini tramite AI specializzata per azienda. Sui PDF testuali dei profili supportati può poi eseguire una verifica deterministica locale e indipendente dall’AI; sulle scansioni Elior applica controlli aritmetici deterministici all’output OCR.", kind="method")

    doc.add_heading("Appendice B · Note di versione", level=1)
    add_table(doc, ["Versione", "Data", "Contenuto"], [
        ["1.2", "19 luglio 2026", "Allineamento al deploy ee9ec2c: stati reali dell’app, separazione per organizzazione, Excel circoscritto a Turni & Riposi e uso mobile aggiornato. Revisione grafica con pagine compatte, schermate incorniciate e tabelle alleggerite."],
        ["1.1", "19 luglio 2026", "Aggiunto il capitolo autonomo sulla Dynamic Island: Live Activity, upload in background, continuità PC–smartphone, strumenti contestuali e limiti."],
        ["1.0", "19 luglio 2026", "Prima edizione completa: acquisizione PC/mobile, estrazione AI, parser deterministico, revisione, archivio, Incidenza, TFR, Turni & Riposi, report e limiti attuali."],
    ], [2.5, 3.5, 11.0])
    add_para(doc, "Il prodotto evolve: prima di usare il manuale in una formazione o allegarlo a una procedura interna, verificare che versione dell’applicazione, profili aziendali, coefficienti e indici coincidano con quelli descritti.", style="Small note")
    add_closing_panel(doc)

    set_headers_and_footers(doc)
    doc.core_properties.title = "ValOra — Manuale operativo"
    doc.core_properties.subject = "Acquisizione, estrazione, controlli deterministici e workflow professionale"
    doc.core_properties.author = "ValOra"
    doc.core_properties.keywords = "ValOra, buste paga, parser deterministico, LegalTech, manuale operativo"
    doc.core_properties.comments = "Generato dal repository ValOra su documentazione tecnica e funzionale verificata."
    return doc


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    doc = build_manual()
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
